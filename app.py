# app.py — Railway Python web server for DOCX reconstruction
# Receives: original DOCX (base64) + rewritten text sections from Claude
# Returns: reconstructed DOCX (base64) with original formatting preserved

import base64
import io
import json
import os
from http.server import HTTPServer, BaseHTTPRequestHandler

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def consolidate_runs(paragraph):
    """Collapse all runs into one, keeping first run's formatting."""
    if not paragraph.runs:
        return
    full_text = paragraph.text
    if not full_text.strip():
        return
    p_el = paragraph._p
    for run in paragraph.runs[1:]:
        p_el.remove(run._r)
    if paragraph.runs:
        paragraph.runs[0].text = full_text


def replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving all formatting."""
    consolidate_runs(paragraph)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = new_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        paragraph._p.append(r)


def normalize(text):
    """Normalize text for matching."""
    import re
    return re.sub(r'\s+', ' ', (text or '').lower().strip())


def get_all_paragraphs(doc):
    """Get all paragraphs including inside tables and text boxes."""
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para not in paragraphs:
                        paragraphs.append(para)
    body = doc.element.body
    for txbx in body.iter(qn('w:txbxContent')):
        for p_el in txbx.iter(qn('w:p')):
            from docx.text.paragraph import Paragraph
            para = Paragraph(p_el, doc)
            if para not in paragraphs:
                paragraphs.append(para)
    return paragraphs


def reconstruct_docx(original_bytes, rewritten_sections):
    """
    Open original DOCX, find each paragraph in rewritten_sections,
    replace text while preserving all formatting XML.
    """
    # Build normalized replacement map
    replacement_map = {}
    for orig, new in rewritten_sections.items():
        replacement_map[normalize(orig)] = new

    doc = Document(io.BytesIO(original_bytes))
    all_paras = get_all_paragraphs(doc)
    replaced_count = 0

    for para in all_paras:
        orig_text = para.text.strip()
        if not orig_text or len(orig_text) < 10:
            continue

        norm = normalize(orig_text)

        # Exact match
        if norm in replacement_map:
            replace_paragraph_text(para, replacement_map[norm])
            replaced_count += 1
            continue

        # Prefix match — first 50 chars
        for key, val in replacement_map.items():
            if len(key) >= 30 and len(norm) >= 30:
                if norm[:50] == key[:50]:
                    replace_paragraph_text(para, val)
                    replaced_count += 1
                    break

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer.read(), replaced_count


class Handler(BaseHTTPRequestHandler):

    def do_POST(self):
        if self.path != '/reconstruct':
            self._json(404, {'error': 'Not found'})
            return

        if not DOCX_AVAILABLE:
            self._json(500, {'error': 'python-docx not installed'})
            return

        content_length = int(self.headers.get('Content-Length', 0))
        body = self.rfile.read(content_length)

        try:
            data = json.loads(body)
        except json.JSONDecodeError:
            self._json(400, {'error': 'Invalid JSON'})
            return

        docx_b64 = data.get('docx_base64')
        rewritten = data.get('rewritten_sections')

        if not docx_b64:
            self._json(400, {'error': 'Missing docx_base64'})
            return
        if not rewritten or not isinstance(rewritten, dict):
            self._json(400, {'error': 'Missing rewritten_sections'})
            return

        try:
            original_bytes = base64.b64decode(docx_b64)
        except Exception:
            self._json(400, {'error': 'Invalid base64'})
            return

        try:
            result_bytes, replaced_count = reconstruct_docx(original_bytes, rewritten)
        except Exception as e:
            self._json(500, {'error': f'Reconstruction failed: {str(e)}'})
            return

        self._json(200, {
            'docx_base64': base64.b64encode(result_bytes).decode('utf-8'),
            'replaced_count': replaced_count,
            'status': 'ok'
        })

    def do_GET(self):
        if self.path == '/health':
            self._json(200, {'status': 'ok', 'docx_available': DOCX_AVAILABLE})
        else:
            self._json(404, {'error': 'Not found'})

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors()
        self.end_headers()

    def _json(self, status, data):
        body = json.dumps(data).encode('utf-8')
        self.send_response(status)
        self._cors()
        self.send_header('Content-Type', 'application/json')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _cors(self):
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')

    def log_message(self, format, *args):
        print(f"[{self.address_string()}] {format % args}")


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    print(f'Starting DOCX reconstruction server on port {port}')
    print(f'python-docx available: {DOCX_AVAILABLE}')
    server = HTTPServer(('0.0.0.0', port), Handler)
    server.serve_forever()
